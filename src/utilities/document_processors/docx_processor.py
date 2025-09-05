import os
import base64
import uuid
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass
import docx
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import zipfile
from PIL import Image
import io
import pandas as pd

from src.schemas import DocumentProcessor, ExtractedContent, ExtractedFigure, ExtractedImage, ExtractedTable

class DOCXProcessor(DocumentProcessor):
    """Enhanced DOCX processor that extracts text, images, tables, and figures using existing schemas"""
    
    def __init__(self):
        self.supported_image_formats = {
            'image/jpeg': '.jpg',
            'image/png': '.png',
            'image/gif': '.gif',
            'image/bmp': '.bmp',
            'image/tiff': '.tiff',
            'image/webp': '.webp'
        }
        self.image_counter = 0
        self.table_counter = 0
        self.figure_counter = 0
    
    def extract_text(self, file_path: str) -> str:
        """Extract plain text from DOCX document"""
        try:
            content = self.extract_complete_content(file_path)
            return content.text
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX {file_path}: {str(e)}")
    
    def extract_complete_content(self, file_path: str) -> ExtractedContent:
        """Extract all content from DOCX document including images, tables, etc."""
        try:
            doc = docx.Document(file_path)
            
            # Reset counters for each document
            self.image_counter = 0
            self.table_counter = 0
            self.figure_counter = 0
            
            # Extract all components
            text = self._extract_formatted_text(doc)
            images = self._extract_images(file_path, doc)
            tables = self._extract_tables(doc)
            figures = self._extract_figures(file_path, doc)
            headers = self._extract_headers(doc)
            footers = self._extract_footers(doc)
            hyperlinks = self._extract_hyperlinks(doc)
            comments = self._extract_comments(doc)
            metadata = self.extract_metadata(file_path)
            
            return ExtractedContent(
                text=text,
                images=images,
                tables=tables,
                figures=figures,
                headers=headers,
                footers=footers,
                hyperlinks=hyperlinks,
                comments=comments,
                metadata=metadata
            )
            
        except Exception as e:
            raise Exception(f"Error extracting complete content from DOCX {file_path}: {str(e)}")
    
    def _extract_formatted_text(self, doc: Document) -> str:
        """Extract text with formatting information preserved"""
        text_parts = []
        
        for element in doc.element.body:
            if isinstance(element, CT_P):
                paragraph = Paragraph(element, doc)
                formatted_text = self._process_paragraph(paragraph)
                if formatted_text.strip():
                    text_parts.append(formatted_text)
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                table_text = self._format_table_as_text(table)
                text_parts.append(table_text)
        
        return "\n\n".join(text_parts)
    
    def _process_paragraph(self, paragraph: Paragraph) -> str:
        """Process paragraph with run-level formatting"""
        if not paragraph.text.strip():
            return ""
        
        formatted_parts = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Add formatting markers
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"_{text}_"
            
            formatted_parts.append(text)
        
        # Handle paragraph style
        result = "".join(formatted_parts)
        
        if paragraph.style.name.startswith('Heading'):
            level = self._get_heading_level(paragraph.style.name)
            result = f"{'#' * level} {result}"
        elif paragraph.style.name == 'List Paragraph':
            result = f"• {result}"
        elif paragraph.style.name in ['Quote', 'Intense Quote']:
            result = f"> {result}"
        
        return result
    
    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name"""
        try:
            if 'Heading' in style_name:
                level_str = style_name.replace('Heading', '').strip()
                return int(level_str) if level_str.isdigit() else 1
        except:
            pass
        return 1
    
    def _extract_images(self, file_path: str, doc: Document) -> List[ExtractedImage]:
        """Extract all images from DOCX document using ExtractedImage schema"""
        images = []
        
        try:
            # Open DOCX as ZIP file to access media
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # Get list of media files
                media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                
                for media_file in media_files:
                    try:
                        self.image_counter += 1
                        
                        # Read image data
                        image_data = docx_zip.read(media_file)
                        
                        # Determine format
                        filename = os.path.basename(media_file)
                        _, ext = os.path.splitext(filename.lower())
                        image_format = ext.lstrip('.')
                        
                        # Get image dimensions and other properties
                        width, height = 0, 0
                        try:
                            with Image.open(io.BytesIO(image_data)) as img:
                                width, height = img.size
                        except:
                            pass
                        
                        # Create ExtractedImage using existing schema
                        extracted_image = ExtractedImage(
                            image_id=f"img_{self.image_counter}_{uuid.uuid4().hex[:8]}",
                            page_number=1,  # DOCX doesn't have page concept like PDF
                            bbox=(0.0, 0.0, float(width), float(height)),  # Approximate bbox
                            image_data=image_data,
                            image_format=image_format,
                            width=width,
                            height=height,
                            file_size=len(image_data),
                            extracted_text=None,  # Could add OCR here if needed
                            description=f"Image from {filename}"
                        )
                        
                        images.append(extracted_image)
                        
                    except Exception as e:
                        print(f"Warning: Could not extract image {media_file}: {str(e)}")
                        continue
                        
        except Exception as e:
            print(f"Warning: Could not extract images from {file_path}: {str(e)}")
        
        return images
    
    def _extract_tables(self, doc: Document) -> List[ExtractedTable]:
        """Extract all tables from the document using ExtractedTable schema"""
        tables = []
        
        for table in doc.tables:
            try:
                self.table_counter += 1
                
                # Extract table data into pandas DataFrame
                data = []
                headers = []
                
                if table.rows:
                    # Get headers from first row
                    first_row = table.rows[0]
                    headers = [cell.text.strip() for cell in first_row.cells]
                    
                    # Get data rows
                    for row in table.rows[1:]:
                        row_data = [cell.text.strip() for cell in row.cells]
                        # Pad row to match header length
                        while len(row_data) < len(headers):
                            row_data.append("")
                        data.append(row_data)
                
                # Create DataFrame
                if headers and data:
                    df = pd.DataFrame(data, columns=headers)
                elif headers:
                    df = pd.DataFrame(columns=headers)
                else:
                    df = pd.DataFrame()
                
                # Generate CSV and HTML representations
                csv_data = df.to_csv(index=False) if not df.empty else ""
                html_data = df.to_html(index=False) if not df.empty else ""
                
                # Create ExtractedTable using existing schema
                extracted_table = ExtractedTable(
                    table_id=f"table_{self.table_counter}_{uuid.uuid4().hex[:8]}",
                    page_number=1,  # DOCX doesn't have page concept
                    bbox=None,  # DOCX doesn't provide precise positioning
                    dataframe=df,
                    confidence=1.0,  # High confidence for native DOCX tables
                    extraction_method="docx_native",
                    csv_data=csv_data,
                    html_data=html_data
                )
                
                tables.append(extracted_table)
                
            except Exception as e:
                print(f"Warning: Could not extract table {self.table_counter}: {str(e)}")
                continue
        
        return tables
    
    def _extract_figures(self, file_path: str, doc: Document) -> List[ExtractedFigure]:
        """Extract figures/charts from DOCX document using ExtractedFigure schema"""
        figures = []
        
        try:
            # In DOCX, figures are often embedded objects or special image types
            # This is a simplified implementation - could be enhanced with shape analysis
            
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # Look for chart data or embedded objects
                chart_files = [f for f in docx_zip.namelist() 
                              if 'chart' in f.lower() or 'embeddings' in f.lower()]
                
                for chart_file in chart_files:
                    try:
                        self.figure_counter += 1
                        
                        # Read figure data
                        figure_data = docx_zip.read(chart_file)
                        
                        # Determine figure type based on file content
                        figure_type = self._determine_figure_type(chart_file)
                        
                        # Create ExtractedFigure using existing schema
                        extracted_figure = ExtractedFigure(
                            figure_id=f"fig_{self.figure_counter}_{uuid.uuid4().hex[:8]}",
                            page_number=1,
                            bbox=(0.0, 0.0, 100.0, 100.0),  # Default bbox
                            figure_type=figure_type,
                            image_data=figure_data,
                            extracted_text=None,
                            caption=f"Figure {self.figure_counter}"
                        )
                        
                        figures.append(extracted_figure)
                        
                    except Exception as e:
                        print(f"Warning: Could not extract figure {chart_file}: {str(e)}")
                        continue
                        
        except Exception as e:
            print(f"Warning: Could not extract figures from {file_path}: {str(e)}")
        
        return figures
    
    def _determine_figure_type(self, filename: str) -> str:
        """Determine figure type from filename"""
        filename_lower = filename.lower()
        if 'chart' in filename_lower:
            return 'chart'
        elif 'graph' in filename_lower:
            return 'graph'
        elif 'diagram' in filename_lower:
            return 'diagram'
        else:
            return 'figure'
    
    def _format_table_as_text(self, table: Table) -> str:
        """Format table as readable text"""
        if not table.rows:
            return ""
        
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        
        if rows:
            # Add separator after header
            separator = " | ".join(["-" * max(len(cell), 3) for cell in rows[0].split(" | ")])
            if len(rows) > 1:
                rows.insert(1, separator)
        
        return "\n".join(rows)
    
    def _extract_headers(self, doc: Document) -> List[str]:
        """Extract header text from all sections"""
        headers = []
        
        for section in doc.sections:
            try:
                if section.header:
                    header_text = ""
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            header_text += paragraph.text.strip() + "\n"
                    if header_text.strip():
                        headers.append(header_text.strip())
            except:
                continue
        
        return headers
    
    def _extract_footers(self, doc: Document) -> List[str]:
        """Extract footer text from all sections"""
        footers = []
        
        for section in doc.sections:
            try:
                if section.footer:
                    footer_text = ""
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            footer_text += paragraph.text.strip() + "\n"
                    if footer_text.strip():
                        footers.append(footer_text.strip())
            except:
                continue
        
        return footers
    
    def _extract_hyperlinks(self, doc: Document) -> List[Dict[str, str]]:
        """Extract all hyperlinks from the document"""
        hyperlinks = []
        
        try:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if hasattr(run._element, 'hyperlink') and run._element.hyperlink is not None:
                        hyperlink = run._element.hyperlink
                        if hasattr(hyperlink, 'address') and hyperlink.address:
                            hyperlinks.append({
                                'text': run.text,
                                'url': hyperlink.address,
                                'type': 'external'
                            })
        except Exception as e:
            print(f"Warning: Could not extract hyperlinks: {str(e)}")
        
        return hyperlinks
    
    def _extract_comments(self, doc: Document) -> List[Dict[str, str]]:
        """Extract all comments from the document"""
        comments = []
        
        try:
            # Note: Full comment extraction would require more complex XML parsing
            # This is a placeholder for future implementation
            pass
        except Exception as e:
            print(f"Warning: Could not extract comments: {str(e)}")
        
        return comments
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract comprehensive metadata from DOCX document"""
        try:
            doc = docx.Document(file_path)
            props = doc.core_properties
            
            # Basic file information
            metadata = {
                'file_path': file_path,
                'file_name': Path(file_path).name,
                'file_type': 'docx',
                'file_size': os.path.getsize(file_path)
            }
            
            # Core properties
            if props.title:
                metadata['title'] = props.title
            if props.author:
                metadata['author'] = props.author
            if props.subject:
                metadata['subject'] = props.subject
            if props.keywords:
                metadata['keywords'] = props.keywords
            if props.comments:
                metadata['comments'] = props.comments
            if props.category:
                metadata['category'] = props.category
            if props.created:
                metadata['created'] = props.created.isoformat()
            if props.modified:
                metadata['modified'] = props.modified.isoformat()
            if props.last_modified_by:
                metadata['last_modified_by'] = props.last_modified_by
            if props.revision:
                metadata['revision'] = props.revision
            
            # Document statistics
            try:
                metadata.update({
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables),
                    'section_count': len(doc.sections),
                })
                
                # Count words approximately
                word_count = 0
                for paragraph in doc.paragraphs:
                    word_count += len(paragraph.text.split())
                metadata['word_count'] = word_count
                
            except Exception as e:
                metadata['stats_error'] = str(e)
            
            return metadata
            
        except Exception as e:
            return {
                'file_path': file_path,
                'file_name': Path(file_path).name,
                'file_type': 'docx',
                'error': str(e)
            }
    
    def extract_structured_data(self, file_path: str) -> Dict[str, Any]:
        """Extract data in a structured format suitable for vector database storage"""
        try:
            content = self.extract_complete_content(file_path)
            
            # Create structured output compatible with existing schemas
            structured_data = {
                'text_content': content.text,
                'metadata': content.metadata,
                'images': [
                    {
                        'image_id': img.image_id,
                        'page_number': img.page_number,
                        'bbox': img.bbox,
                        'image_format': img.image_format,
                        'width': img.width,
                        'height': img.height,
                        'file_size': img.file_size,
                        'description': img.description,
                        'base64_data': base64.b64encode(img.image_data).decode('utf-8')
                    } for img in content.images
                ],
                'tables': [
                    {
                        'table_id': table.table_id,
                        'page_number': table.page_number,
                        'confidence': table.confidence,
                        'extraction_method': table.extraction_method,
                        'csv_data': table.csv_data,
                        'html_data': table.html_data,
                        'dataframe_info': {
                            'shape': table.dataframe.shape,
                            'columns': table.dataframe.columns.tolist()
                        }
                    } for table in content.tables
                ],
                'figures': [
                    {
                        'figure_id': fig.figure_id,
                        'page_number': fig.page_number,
                        'bbox': fig.bbox,
                        'figure_type': fig.figure_type,
                        'caption': fig.caption,
                        'base64_data': base64.b64encode(fig.image_data).decode('utf-8')
                    } for fig in content.figures
                ],
                'headers': content.headers,
                'footers': content.footers,
                'hyperlinks': content.hyperlinks,
                'comments': content.comments
            }
            
            return structured_data
            
        except Exception as e:
            raise Exception(f"Error extracting structured data from {file_path}: {str(e)}")

# Enhanced utility functions for vector database integration
def process_docx_for_vector_storage(file_path: str, 
                                   chunk_tables_separately: bool = True,
                                   chunk_images_separately: bool = True,
                                   chunk_figures_separately: bool = True) -> List[Dict[str, Any]]:
    """
    Process DOCX file and return chunks suitable for vector database storage
    Compatible with existing PDF processing pipeline
    """
    processor = DOCXProcessor()
    structured_data = processor.extract_structured_data(file_path)
    
    chunks = []
    
    # Main text content chunk
    if structured_data['text_content'].strip():
        chunks.append({
            'content': structured_data['text_content'],
            'content_type': 'text',
            'metadata': structured_data['metadata'],
            'source': file_path,
            'extraction_method': 'docx_native'
        })
    
    # Table chunks (if requested)
    if chunk_tables_separately:
        for table in structured_data['tables']:
            table_content = f"Table ID: {table['table_id']}\n"
            table_content += f"Extraction Method: {table['extraction_method']}\n"
            table_content += f"Confidence: {table['confidence']}\n\n"
            table_content += "CSV Data:\n" + table['csv_data']
            
            chunks.append({
                'content': table_content,
                'content_type': 'table',
                'metadata': {
                    **structured_data['metadata'],
                    'table_id': table['table_id'],
                    'extraction_method': table['extraction_method'],
                    'confidence': table['confidence'],
                    'table_shape': table['dataframe_info']['shape']
                },
                'source': file_path,
                'csv_data': table['csv_data'],
                'html_data': table['html_data']
            })
    
    # Image chunks (if requested)
    if chunk_images_separately:
        for img in structured_data['images']:
            image_content = f"Image ID: {img['image_id']}\n"
            image_content += f"Description: {img['description']}\n"
            image_content += f"Dimensions: {img['width']}x{img['height']}\n"
            image_content += f"Format: {img['image_format']}\n"
            image_content += f"File Size: {img['file_size']} bytes"
            
            chunks.append({
                'content': image_content,
                'content_type': 'image',
                'metadata': {
                    **structured_data['metadata'],
                    'image_id': img['image_id'],
                    'image_format': img['image_format'],
                    'image_dimensions': (img['width'], img['height']),
                    'file_size': img['file_size']
                },
                'image_data': img['base64_data'],
                'source': file_path
            })
    
    # Figure chunks (if requested)
    if chunk_figures_separately:
        for fig in structured_data['figures']:
            figure_content = f"Figure ID: {fig['figure_id']}\n"
            figure_content += f"Type: {fig['figure_type']}\n"
            if fig['caption']:
                figure_content += f"Caption: {fig['caption']}\n"
            
            chunks.append({
                'content': figure_content,
                'content_type': 'figure',
                'metadata': {
                    **structured_data['metadata'],
                    'figure_id': fig['figure_id'],
                    'figure_type': fig['figure_type'],
                    'caption': fig['caption']
                },
                'image_data': fig['base64_data'],
                'source': file_path
            })
    
    return chunks

# Integration with existing vector database system
def add_docx_to_vector_db(file_path: str, vector_db, embedding_function,
                         chunk_tables_separately: bool = True,
                         chunk_images_separately: bool = True,
                         chunk_figures_separately: bool = True):
    """
    Add processed DOCX content to vector database using existing schemas
    """
    chunks = process_docx_for_vector_storage(
        file_path, 
        chunk_tables_separately,
        chunk_images_separately,
        chunk_figures_separately
    )
    
    documents = []
    for chunk in chunks:
        # Generate embedding for the content
        embedding = embedding_function(chunk['content'])
        
        document = {
            'text': chunk['content'],
            'embedding': embedding,
            'metadata': chunk['metadata'],
            'content_type': chunk['content_type'],
            'source': chunk['source']
        }
        
        # Add additional data based on content type
        if 'image_data' in chunk:
            document['image_data'] = chunk['image_data']
        
        if 'csv_data' in chunk:
            document['csv_data'] = chunk['csv_data']
            document['html_data'] = chunk['html_data']
        
        documents.append(document)
    
    # Add to vector database
    doc_ids = vector_db.add_documents(documents)
    return doc_ids